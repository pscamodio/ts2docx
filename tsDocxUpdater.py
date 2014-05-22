from docx import Document
from pyqtts import tsfile
from os import path
import argparse

parser = argparse.ArgumentParser(description='Create docx file from ts file')
parser.add_argument('input', type=str, help='input ts file')
parser.add_argument('--docx','-d', dest='docx', type=str,
                    help='input docx file, if missing searching a filename similar to input')
parser.add_argument('--output', '-o', dest='output', type=str, help='output ts file, if missing overwrite input')


def get_text_from_cell(cell):
    return "\n".join([p.text or "" for p in cell.paragraphs])


def my_strip(s):
    if s:
        return "".join(s.split())
    else:
        return ""


def main():
    args = parser.parse_args()

    if not args.docx:
        args.docx = path.splitext(args.input)[0] + ".docx"
    if not args.output:
        args.output = args.input

    print("Opening ts file: " + args.input)
    orig_ts = tsfile.tsfile()
    orig_ts.load_file(args.input)

    print("Opening docx file: " + args.docx)
    doc = Document(args.docx)
    new_table = doc.tables[0]
    rows = [row.cells for row in new_table.rows]
    new_data = {}
    print("Reading translation from docx")
    for row in rows:
        context = get_text_from_cell(row[0] or "")
        orig = get_text_from_cell(row[1] or "")
        transl = get_text_from_cell(row[2] or "")
        new_data[(my_strip(context), my_strip(orig))] = transl
    print("Found " + str(len(new_data)) + " translations")
    print("Updating ts data")
    for context in orig_ts.context_list:
        for message in context.message_list:
            try:
                message.translation.text = new_data[(my_strip(context.name),my_strip(message.source))]
                message.translation.finished = message.translation.text != ""
            except KeyError:
                pass

    print("Saving new tsdata in: " + args.output)
    orig_ts.save_file(args.output)



if __name__ == "__main__":
    main()